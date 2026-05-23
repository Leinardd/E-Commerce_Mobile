"""
Generate SRS .docx for Varón E-Commerce Mobile Application
Run: python generate_srs.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Helper functions ──────────────────────────────────────────────────────────

def set_font(run, name='Times New Roman', size=12, bold=False,
             italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(text='', style='Normal', align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6, bold=False, italic=False,
                  size=12, font='Times New Roman'):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, font, size, bold, italic)
    return p

def add_heading(text, level=1, space_before=12, space_after=6):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, 'Times New Roman', sizes.get(level, 12),
             bold=True, color=(0, 0, 0))
    return p

def add_bullet(text, level=0, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    set_font(run, 'Times New Roman', size)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, 'Times New Roman', 11, bold=True)
        cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'D9D9D9')
        shd.set(qn('w:val'), 'clear')
        cell._tc.tcPr.append(shd)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_font(run, 'Times New Roman', 11)
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
run = p.add_run('SOFTWARE REQUIREMENT SPECIFICATION')
set_font(run, 'Times New Roman', 18, bold=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Varón: A Multi-Vendor E-Commerce Mobile Application\nfor the Philippine Fashion Retail Market"')
set_font(run, 'Times New Roman', 14, bold=True, italic=True)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

for label, value in [
    ('May 2026', None),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    set_font(run, 'Times New Roman', 12)

doc.add_paragraph()

add_paragraph('Prepared by:', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
for m in ['(Member Name 1)', '(Member Name 2)', '(Member Name 3)']:
    add_paragraph(f'• {m}', align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

fields = [
    ('Course:', 'Bachelor of Science in Information Technology'),
    ('Subject:', 'ITEC-306 Application Development and Emerging Technologies'),
    ('Submitted to:', '(Instructor Name)'),
    ('Institution:', '(School/University Name)'),
    ('Date Submitted:', '(Month Day, Year)'),
]
for label, value in fields:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label} ')
    set_font(r1, 'Times New Roman', 12, bold=True)
    r2 = p.add_run(value)
    set_font(r2, 'Times New Roman', 12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════

add_heading('TABLE OF CONTENTS', 1, space_before=0)
doc.add_paragraph()

toc_entries = [
    ('1. Introduction', '3'),
    ('   1.1 Purpose', '3'),
    ('   1.2 Scope', '3'),
    ('   1.3 Definitions, Acronyms, and Abbreviations', '4'),
    ('   1.4 References', '4'),
    ('   1.5 Overview', '4'),
    ('2. Overall Description', '5'),
    ('   2.1 Product Perspective', '5'),
    ('   2.2 Product Functions', '5'),
    ('   2.3 User Classes and Characteristics', '6'),
    ('   2.4 Operating Environment', '7'),
    ('   2.5 Design and Implementation Constraints', '7'),
    ('   2.6 Assumptions and Dependencies', '7'),
    ('3. System Features and Functional Requirements', '8'),
    ('   3.1 User Authentication and Role Management', '8'),
    ('   3.2 Buyer – Product Browsing and Search', '9'),
    ('   3.3 Buyer – Shopping Cart and Checkout', '10'),
    ('   3.4 Buyer – Order Tracking', '11'),
    ('   3.5 Seller – Store and Product Management', '11'),
    ('   3.6 Seller – Order Management', '13'),
    ('   3.7 Rider – Delivery Management', '13'),
    ('   3.8 Admin – Platform Management', '14'),
    ('   3.9 Notifications', '15'),
    ('4. External Interface Requirements', '15'),
    ('   4.1 User Interfaces', '15'),
    ('   4.2 Hardware Interfaces', '16'),
    ('   4.3 Software Interfaces', '16'),
    ('   4.4 Communication Interfaces', '16'),
    ('5. Non-Functional Requirements', '17'),
    ('   5.1 Performance Requirements', '17'),
    ('   5.2 Security Requirements', '17'),
    ('   5.3 Usability Requirements', '17'),
    ('   5.4 Reliability and Availability', '18'),
    ('   5.5 Maintainability and Scalability', '18'),
    ('6. Other Requirements', '18'),
    ('   6.1 Database Requirements', '18'),
    ('   6.2 Legal and Compliance Requirements', '19'),
]
for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(entry)
    set_font(run, 'Times New Roman', 11)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

add_heading('1. Introduction', 1)

add_heading('1.1 Purpose', 2)
add_paragraph(
    'This Software Requirements Specification (SRS) document describes the functional and '
    'non-functional requirements for the Varón E-Commerce Mobile Application. Varón is a '
    'multi-vendor online marketplace designed for the Philippine fashion retail market, '
    'enabling buyers to browse and purchase fashion products, sellers to manage their '
    'online stores, and riders to fulfill deliveries. This document is intended for use by '
    'the development team, project stakeholders, course instructors, and evaluators of '
    'ITEC-306: Application Development and Emerging Technologies. It serves as the '
    'authoritative reference for the system\'s design, development, and validation.'
)

add_heading('1.2 Scope', 2)
add_paragraph(
    'The Varón E-Commerce Mobile Application (hereafter referred to as "the System" or '
    '"Varón") is a cross-platform mobile application developed using the Flutter framework '
    'with Supabase as the backend-as-a-service platform. The System supports four distinct '
    'user roles: Buyer, Seller, Rider, and Administrator.'
)
add_paragraph(
    'The System provides the following high-level capabilities:'
)
add_bullet('A product marketplace where buyers can browse, search, and purchase fashion items from multiple sellers.')
add_bullet('A seller portal where registered and approved sellers can manage their store, product listings, inventory, and orders.')
add_bullet('A rider management module where approved delivery riders can accept, track, and fulfill deliveries.')
add_bullet('An administrator dashboard (web-only) for platform governance, including user management, product approval, and order oversight.')
add_bullet('Role-based authentication and authorization using Supabase Authentication with email and OTP verification.')
doc.add_paragraph()
add_paragraph(
    'The System does not include an integrated payment gateway in its current version. Orders '
    'are placed with a "To Pay" status, and payment processing is designated for a future '
    'development phase. The administrator portal is accessible via web browser only and is '
    'intentionally restricted on mobile devices.'
)

add_heading('1.3 Definitions, Acronyms, and Abbreviations', 2)
add_table(
    ['Term / Acronym', 'Definition'],
    [
        ('SRS', 'Software Requirements Specification'),
        ('UI', 'User Interface'),
        ('API', 'Application Programming Interface'),
        ('OTP', 'One-Time Password — a six-digit code sent to a user\'s email for verification'),
        ('RLS', 'Row-Level Security — Supabase/PostgreSQL policy controlling data access per user'),
        ('CRUD', 'Create, Read, Update, Delete — standard data operations'),
        ('SKU', 'Stock Keeping Unit — unique identifier for a product variant'),
        ('FK', 'Foreign Key — a database column referencing another table\'s primary key'),
        ('UUID', 'Universally Unique Identifier — used by Supabase Auth for user identification'),
        ('Buyer', 'An end-user who browses and purchases products on the platform'),
        ('Seller', 'A registered merchant who lists and sells products through the platform'),
        ('Rider', 'A registered delivery personnel who accepts and fulfills delivery orders'),
        ('Admin', 'A platform administrator who manages users, products, and orders'),
        ('Varón', 'The name of the e-commerce mobile application (Spanish for "male/man")'),
        ('Flutter', 'An open-source UI software development kit by Google for cross-platform apps'),
        ('Supabase', 'An open-source backend-as-a-service platform providing PostgreSQL, Auth, and Storage'),
        ('BSIT', 'Bachelor of Science in Information Technology'),
    ],
    [1.5, 4.5]
)

add_heading('1.4 References', 2)
add_bullet('IEEE Std 830-1998: IEEE Recommended Practice for Software Requirements Specifications.')
add_bullet('Supabase Documentation. Retrieved from https://supabase.com/docs')
add_bullet('Flutter Documentation. Retrieved from https://flutter.dev/docs')
add_bullet('Google Fonts – Inter & Commissioner typefaces used in the application UI.')
add_bullet('flutter_map package documentation. Retrieved from https://pub.dev/packages/flutter_map')
add_bullet('geolocator package documentation. Retrieved from https://pub.dev/packages/geolocator')
doc.add_paragraph()

add_heading('1.5 Overview', 2)
add_paragraph(
    'The remainder of this document is organized as follows:'
)
add_bullet('Section 2 provides an overall description of the product, including its context, major functions, user characteristics, operating environment, and constraints.')
add_bullet('Section 3 details the specific system features and functional requirements, organized by user role.')
add_bullet('Section 4 describes external interface requirements, including user, hardware, software, and communication interfaces.')
add_bullet('Section 5 presents non-functional requirements covering performance, security, usability, reliability, and maintainability.')
add_bullet('Section 6 covers additional database and legal requirements.')
doc.add_paragraph()

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — OVERALL DESCRIPTION
# ══════════════════════════════════════════════════════════════════════════════

add_heading('2. Overall Description', 1)

add_heading('2.1 Product Perspective', 2)
add_paragraph(
    'Varón is a standalone mobile application that operates as a multi-vendor e-commerce '
    'marketplace focused on the Philippine men\'s fashion retail sector. The System is '
    'self-contained and communicates exclusively with Supabase cloud services for data '
    'persistence, authentication, and file storage. It is not a replacement or extension of '
    'any existing proprietary system.'
)
add_paragraph(
    'The System architecture follows a client-server model:'
)
add_bullet('Client: Flutter mobile application (Android and iOS) and web browser (for the Admin portal).')
add_bullet('Server: Supabase project hosting a PostgreSQL database, Supabase Auth (JWT-based), and Supabase Storage for product images.')
add_bullet('Maps and Location: flutter_map widget powered by OpenStreetMap tiles; device GPS accessed via the geolocator package.')
doc.add_paragraph()
add_paragraph(
    'Local state is managed through SharedPreferences for session persistence and cart '
    'management. The shopping cart is maintained as an in-memory singleton during an active '
    'session.'
)

add_heading('2.2 Product Functions', 2)
add_paragraph('The System provides the following major functions, grouped by user role:')

add_paragraph('Buyer Functions:', bold=True, space_after=2)
add_bullet('Register and authenticate using email and password with optional OTP email verification.')
add_bullet('Browse products by category (Tops, Bottoms, Footwear, Accessories, Grooming, Barong, Suits, etc.).')
add_bullet('View product details including images, description, price, stock, size guide, and seller information.')
add_bullet('Select product variants (size and color) and add items to a shopping cart.')
add_bullet('Proceed to checkout, enter a delivery address, and place an order.')
add_bullet('Track order status across multiple stages: To Pay → To Ship → Shipped → To Receive → Completed.')
add_bullet('View purchase history and order details.')
add_bullet('Navigate to a seller\'s storefront and browse all products from that seller.')

add_paragraph('Seller Functions:', bold=True, space_after=2)
add_bullet('Register as a seller by submitting an application (store name, full name, address, phone).')
add_bullet('Await admin approval before gaining access to the seller portal.')
add_bullet('Create, edit, and delete product listings with images, variants, pricing, stock, and delivery options.')
add_bullet('Monitor product approval status (Pending, Approved, Rejected) and view rejection reasons.')
add_bullet('Receive and manage buyer orders; update order status from To Pay to To Ship.')
add_bullet('View dashboard analytics: revenue, total products, active products, and recent orders.')
add_bullet('Receive in-app notifications for product approval and rejection events.')

add_paragraph('Rider Functions:', bold=True, space_after=2)
add_bullet('Register as a delivery rider by submitting an application (license number, vehicle type).')
add_bullet('Await admin approval before accessing the rider portal.')
add_bullet('Browse available orders ready for pickup (status: To Ship, unassigned).')
add_bullet('Accept delivery orders and progress them through delivery stages: Rider Accepted → Picked Up → In Transit → Near Location → Delivered.')
add_bullet('View live delivery map with GPS tracking.')
add_bullet('Track completed deliveries and accumulated earnings (15% commission per delivered order).')

add_paragraph('Admin Functions (Web-only):', bold=True, space_after=2)
add_bullet('View platform-wide statistics: total users, sellers, riders, orders, and revenue.')
add_bullet('Manage users: view profiles, ban, or activate accounts.')
add_bullet('Approve or suspend seller and rider applications.')
add_bullet('Approve or reject product listings submitted by sellers, with optional rejection reason.')
add_bullet('View and manage all platform orders.')
doc.add_paragraph()

add_heading('2.3 User Classes and Characteristics', 2)
add_table(
    ['User Class', 'Description', 'Technical Proficiency', 'Access Level'],
    [
        ('Buyer',
         'End-consumers who browse, add to cart, and purchase fashion products. Primarily mobile users aged 18–45.',
         'Low to moderate. Expected to use basic mobile app interactions.',
         'Public product catalog, own orders and profile.'),
        ('Seller',
         'Individual merchants or small business owners who list and sell fashion items. Must be approved by Admin.',
         'Moderate. Expected to manage product listings, images, and orders.',
         'Own store, product listings, and related orders.'),
        ('Rider',
         'Delivery personnel who accept and fulfill order deliveries. Must be approved by Admin.',
         'Low to moderate. Expected to accept orders and update delivery statuses.',
         'Available orders (unassigned), own active deliveries and earnings.'),
        ('Administrator',
         'Platform operators responsible for governance and oversight. Web-only access.',
         'High. Manages platform-level operations including user and product approval.',
         'Full platform data (users, sellers, riders, orders, products).'),
    ],
    [1.2, 2.8, 1.5, 2.0]
)

add_heading('2.4 Operating Environment', 2)
add_paragraph(
    'The System operates in the following environments:'
)
add_bullet('Mobile Platform: Android (API Level 21 and above) and iOS (iOS 12 and above).')
add_bullet('Web Platform: Admin portal accessible on modern web browsers (Chrome, Firefox, Edge) via Flutter Web compilation.')
add_bullet('Backend: Supabase cloud services hosted on AWS infrastructure.')
add_bullet('Network: Requires an active internet connection for all data operations, authentication, and image loading.')
add_bullet('Storage: Supabase Storage ("products" bucket) for product image hosting.')
add_bullet('Maps: OpenStreetMap tiles via flutter_map; device GPS via geolocator.')
doc.add_paragraph()

add_heading('2.5 Design and Implementation Constraints', 2)
add_bullet('The application is developed using Flutter (Dart), requiring Dart SDK 3.x and Flutter SDK 3.x or higher.')
add_bullet('All backend services use Supabase (PostgreSQL). No alternative database or backend is supported in this version.')
add_bullet('The shopping cart is stored in memory and is not persisted across app restarts in the current version.')
add_bullet('No payment gateway is integrated. Orders remain in "To Pay" status until future payment integration.')
add_bullet('The Admin dashboard is intentionally blocked on mobile devices; it is accessible only via web browser.')
add_bullet('Product images are uploaded to Supabase Storage and are publicly accessible via generated URLs.')
add_bullet('Row-Level Security (RLS) policies in Supabase govern all data access to ensure users can only access authorized records.')
add_bullet('The application requires a Google Maps/Geolocator-compatible device for rider map functionality.')
doc.add_paragraph()

add_heading('2.6 Assumptions and Dependencies', 2)
add_bullet('It is assumed that Supabase project credentials (URL and anon key) are securely configured in the application environment.')
add_bullet('It is assumed that sellers and riders will have valid identification and contact information at the time of application.')
add_bullet('The System depends on Supabase Auth for all authentication. If Supabase Auth is unavailable, login and signup will fail.')
add_bullet('The System depends on device GPS hardware and user permission grants for the rider map feature.')
add_bullet('It is assumed that the admin will actively monitor and process seller and rider applications in a timely manner.')
add_bullet('Email delivery for OTP verification depends on Supabase\'s SMTP configuration.')
add_bullet('The application assumes a stable internet connection is available during checkout and order operations.')
doc.add_paragraph()

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading('3. System Features and Functional Requirements', 1)
add_paragraph(
    'This section describes the functional requirements of the Varón System organized by '
    'feature area. Each requirement is presented with a unique identifier (FR-XXX), a '
    'description, and its priority level (High / Medium / Low).'
)
doc.add_paragraph()

# 3.1
add_heading('3.1 User Authentication and Role Management', 2)
add_paragraph(
    'This feature covers all authentication, registration, and role-based access control '
    'capabilities of the System.'
)
add_table(
    ['ID', 'Requirement Description', 'Priority'],
    [
        ('FR-101', 'The System shall allow buyers to register using a valid email address, password (minimum 6 characters), first name, last name, and phone number.', 'High'),
        ('FR-102', 'The System shall support email OTP (One-Time Password) verification during buyer registration when email confirmation is enabled in Supabase Auth.', 'High'),
        ('FR-103', 'The System shall allow buyers, sellers, and riders to log in using their registered email address and password via Supabase Authentication.', 'High'),
        ('FR-104', 'The System shall support password reset via a 6-digit OTP code sent to the user\'s registered email address.', 'High'),
        ('FR-105', 'The System shall determine each authenticated user\'s role (buyer, seller, rider, admin) from the users table and redirect them to the appropriate portal.', 'High'),
        ('FR-106', 'The System shall allow buyers to apply as sellers by submitting a form with store name, full name, address, and phone number.', 'High'),
        ('FR-107', 'The System shall allow buyers to apply as riders by submitting a form with driver\'s license number, vehicle type, and address.', 'High'),
        ('FR-108', 'The System shall restrict Admin portal access to web browsers and display an "Unauthorized" screen when an admin attempts mobile access.', 'High'),
        ('FR-109', 'The System shall support logout for all roles, clearing the local session and signing out from Supabase Auth.', 'High'),
        ('FR-110', 'The System shall display a "Pending Approval" screen for seller and rider accounts whose status is "pending" at login.', 'Medium'),
        ('FR-111', 'The System shall automatically create a user profile in the users table upon successful authentication if one does not yet exist.', 'Medium'),
        ('FR-112', 'The System shall prevent banned accounts from logging in and display an appropriate message.', 'High'),
    ],
    [0.7, 5.2, 0.7]
)

# 3.2
add_heading('3.2 Buyer – Product Browsing and Search', 2)
add_paragraph(
    'This feature covers how buyers discover and explore products within the Varón marketplace.'
)
add_table(
    ['ID', 'Requirement Description', 'Priority'],
    [
        ('FR-201', 'The System shall display a landing home page showing a curated list of products, including a featured section of the four most recently approved products.', 'High'),
        ('FR-202', 'The System shall allow buyers to browse products organized by categories (e.g., Tops, Bottoms, Footwear, Accessories, Grooming, Barong, Suits & Blazers).', 'High'),
        ('FR-203', 'The System shall display only products with approval_status = "approved" and archive_status = "active" to buyers.', 'High'),
        ('FR-204', 'The System shall display a product card showing the product image, name, seller name, price, and an "Add to Cart" button.', 'High'),
        ('FR-205', 'The System shall navigate buyers to a Product Detail Screen when a product card is tapped.', 'High'),
        ('FR-206', 'The Product Detail Screen shall display the product name, category, price, description, stock status, image gallery (thumbnail strip), available colors, available sizes, and size guide.', 'High'),
        ('FR-207', 'The System shall display a size guide (chart with measurements) for applicable categories: Tops, Bottoms, and Footwear.', 'Medium'),
        ('FR-208', 'The System shall display the seller\'s store name on the product card as a clickable link that navigates to the Seller Store Screen.', 'High'),
        ('FR-209', 'The System shall display an "About the Seller" card on the Product Detail Screen showing the seller\'s store name, logo, and a "Visit Shop" button.', 'High'),
        ('FR-210', 'The Seller Store Screen shall display the seller\'s store name, logo, banner, description, product count, category filters, and a grid of all the seller\'s approved products.', 'High'),
        ('FR-211', 'The System shall display a stock indicator badge on Product Detail Screen (e.g., "In Stock", "Only X left", "Out of Stock").', 'Medium'),
        ('FR-212', 'The System shall reload the product list when a Supabase authentication session is established, ensuring products appear even on first launch without a prior login.', 'High'),
    ],
    [0.7, 5.2, 0.7]
)

# 3.3
add_heading('3.3 Buyer – Shopping Cart and Checkout', 2)
add_table(
    ['ID', 'Requirement Description', 'Priority'],
    [
        ('FR-301', 'The System shall allow buyers to add products to a shopping cart by selecting a size and color variant through the Variant Picker Sheet.', 'High'),
        ('FR-302', 'The System shall maintain a unique cart entry per product-size-color combination (variant key), preventing duplicate variant entries.', 'High'),
        ('FR-303', 'The System shall allow buyers to adjust item quantities within the cart.', 'High'),
        ('FR-304', 'The System shall allow buyers to remove items from the cart.', 'High'),
        ('FR-305', 'The System shall display the cart total price, itemized list, and selected variants on the Cart Screen.', 'High'),
        ('FR-306', 'The System shall prevent adding out-of-stock products to the cart (the "Add to Cart" button shall be disabled when stock = 0).', 'High'),
        ('FR-307', 'The Checkout Screen shall require buyers to enter a complete delivery address including house/unit number, street, barangay, city/municipality, province, and region.', 'High'),
        ('FR-308', 'The System shall display an order summary on the Checkout Screen showing product name, variant, quantity, unit price, and total amount.', 'High'),
        ('FR-309', 'Upon order placement, the System shall create one order record per product in the orders data structure with status "toPay".', 'High'),
        ('FR-310', 'The System shall clear the cart after a successful order placement.', 'High'),
        ('FR-311', 'The System shall support a "Buy Now" flow that adds an item to the cart and proceeds directly to Checkout without requiring the buyer to visit the Cart Screen.', 'Medium'),
    ],
    [0.7, 5.2, 0.7]
)

# 3.4
add_heading('3.4 Buyer – Order Tracking', 2)
add_table(
    ['ID', 'Requirement Description', 'Priority'],
    [
        ('FR-401', 'The System shall display a buyer\'s order history on the Buyer Orders Screen, showing all past and active orders.', 'High'),
        ('FR-402', 'The System shall display the current status of each order using the following progression: To Pay → To Ship → Shipped → To Receive → Completed / Cancelled.', 'High'),
        ('FR-403', 'The System shall display order details including product name, image, variant, quantity, price, delivery address, and order date.', 'High'),
        ('FR-404', 'The System shall update the order status display in real time when the seller or rider updates it.', 'Medium'),
    ],
    [0.7, 5.2, 0.7]
)

# 3.5
add_heading('3.5 Seller – Store and Product Management', 2)
add_table(
    ['ID', 'Requirement Description', 'Priority'],
    [
        ('FR-501', 'The System shall display a Seller Dashboard showing key metrics: total products, active products, products pending approval, archived products, low-stock products (< 10 units), total orders, total revenue, and the top-selling product.', 'High'),
        ('FR-502', 'The System shall allow sellers to add new products with the following information: name, description, price, stock quantity, category, delivery options (delivery/pickup/both), condition (new/refurbished/used), and at least one product image.', 'High'),
        ('FR-503', 'The System shall allow sellers to upload multiple images per product to Supabase Storage.', 'High'),
        ('FR-504', 'The System shall allow sellers to define product variants (color and size combinations) with individual stock levels.', 'High'),
        ('FR-505', 'Newly created products shall be assigned approval_status = "pending" and will not be visible to buyers until approved by an Admin.', 'High'),
        ('FR-506', 'The System shall allow sellers to view their product list filtered by approval status (All, Pending, Approved, Rejected) and stock level.', 'High'),
        ('FR-507', 'The System shall allow sellers to edit existing product information, including name, description, price, stock, images, and variants.', 'High'),
        ('FR-508', 'The System shall allow sellers to archive and unarchive products, making them hidden from the buyer marketplace without deletion.', 'Medium'),
        ('FR-509', 'The System shall allow sellers to delete products that are no longer needed.', 'Medium'),
        ('FR-510', 'The System shall display rejection reasons for products that have been rejected by the Admin.', 'High'),
        ('FR-511', 'The Seller Dashboard shall display a revenue chart showing daily revenue for the past seven days.', 'Medium'),
        ('FR-512', 'The System shall display a public Seller Store Screen accessible to buyers, showing the seller\'s store name, logo, banner, description, product count, and a filterable product grid.', 'High'),
    ],
    [0.7, 5.2, 0.7]
)

# 3.6
add_heading('3.6 Seller – Order Management', 2)
add_table(
    ['ID', 'Requirement Description', 'Priority'],
    [
        ('FR-601', 'The System shall display all orders associated with the seller\'s products on the Seller Orders Screen.', 'High'),
        ('FR-602', 'The Seller Orders Screen shall support filtering orders by status (All, To Pay, To Ship, Shipped, Completed, Cancelled).', 'Medium'),
        ('FR-603', 'The System shall allow sellers to update order status from "To Pay" to "To Ship", indicating the item is packaged and ready for rider pickup.', 'High'),
        ('FR-604', 'The System shall display buyer delivery address, product name, quantity, and variant for each order visible to the seller.', 'High'),
    ],
    [0.7, 5.2, 0.7]
)

# 3.7
add_heading('3.7 Rider – Delivery Management', 2)
add_table(
    ['ID', 'Requirement Description', 'Priority'],
    [
        ('FR-701', 'The System shall display a Rider Dashboard showing: available orders count, active deliveries count, completed deliveries count, and total earnings.', 'High'),
        ('FR-702', 'The System shall display a list of available orders (status = "To Ship", no rider assigned) on the Rider Available Orders Screen.', 'High'),
        ('FR-703', 'The System shall allow riders to accept an available order, setting the order status to "Rider Accepted" and assigning the rider\'s identifier.', 'High'),
        ('FR-704', 'The System shall allow riders to progress delivery status through the following stages: Rider Accepted → Picked Up → In Transit → Near Location → Delivered.', 'High'),
        ('FR-705', 'The System shall prevent other riders from accepting an order that has already been accepted.', 'High'),
        ('FR-706', 'The Rider Map Screen shall display the rider\'s current GPS location on an interactive map powered by flutter_map and OpenStreetMap.', 'High'),
        ('FR-707', 'The System shall calculate and display rider earnings as 15% of the completed order\'s total amount.', 'High'),
        ('FR-708', 'The Rider Earnings Screen shall display a history of completed deliveries and the corresponding earnings, including a 7-day earnings summary.', 'Medium'),
    ],
    [0.7, 5.2, 0.7]
)

# 3.8
add_heading('3.8 Admin – Platform Management', 2)
add_paragraph(
    'The Admin portal is accessible exclusively via web browser. Mobile access is blocked '
    'and redirected to the Unauthorized Screen.'
)
add_table(
    ['ID', 'Requirement Description', 'Priority'],
    [
        ('FR-801', 'The Admin Dashboard shall display platform-wide statistics including total registered users, buyers, sellers, riders, orders, products, and total revenue.', 'High'),
        ('FR-802', 'The System shall allow admins to view a list of all registered users with their role, account status, and registration date.', 'High'),
        ('FR-803', 'The System shall allow admins to ban or activate user accounts.', 'High'),
        ('FR-804', 'The System shall allow admins to view all pending seller applications and approve or suspend them.', 'High'),
        ('FR-805', 'The System shall allow admins to view all pending rider applications and approve or suspend them.', 'High'),
        ('FR-806', 'The System shall allow admins to view all pending product listings and approve them, making them visible to buyers.', 'High'),
        ('FR-807', 'The System shall allow admins to reject product listings with a mandatory rejection reason that will be displayed to the seller.', 'High'),
        ('FR-808', 'The System shall allow admins to view all platform orders filtered by status.', 'Medium'),
    ],
    [0.7, 5.2, 0.7]
)

# 3.9
add_heading('3.9 Notifications', 2)
add_table(
    ['ID', 'Requirement Description', 'Priority'],
    [
        ('FR-901', 'The System shall generate an in-app notification for a seller when their product is approved by the Admin.', 'High'),
        ('FR-902', 'The System shall generate an in-app notification for a seller when their product is rejected by the Admin, including the rejection reason.', 'High'),
        ('FR-903', 'The Seller Notifications Screen shall display all approval and rejection notifications with read/unread status.', 'High'),
        ('FR-904', 'The System shall display an unread notification count badge on the Seller Dashboard to indicate new unread notifications.', 'Medium'),
    ],
    [0.7, 5.2, 0.7]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — EXTERNAL INTERFACE REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading('4. External Interface Requirements', 1)

add_heading('4.1 User Interfaces', 2)
add_paragraph(
    'The Varón application provides the following distinct user interface areas:'
)
add_bullet('Landing Screen: Entry point displaying product listings, category filters, and a navigation bar for unauthenticated and authenticated buyers.')
add_bullet('Authentication Screens: A unified Login/Signup screen with toggle between login and registration forms, email OTP verification screen, and forgot password screen with OTP reset flow.')
add_bullet('Buyer Portal: Bottom navigation bar with Home, Shop, and Profile tabs. Product detail screens, cart screen, checkout screen, and order tracking screen.')
add_bullet('Seller Portal: Dashboard with statistics and quick actions. Product management screens with image upload and variant configuration. Order management and notification screens.')
add_bullet('Rider Portal: Dashboard with stats and quick actions. Available orders list, delivery detail screen, live map screen, and earnings history screen.')
add_bullet('Admin Portal (Web): Four-tab dashboard (Overview, Users, Orders, Products) accessible via web browser only.')
doc.add_paragraph()
add_paragraph(
    'The UI is designed following a minimalist, monochromatic aesthetic consistent with '
    'premium fashion branding. Typography uses Google Fonts (Inter for body text, '
    'Commissioner for headings and labels, Playfair Display for product names). The primary '
    'color scheme uses near-black (#0A0A0A), white (#FFFFFF), and light grey (#EEEEEE) tones. '
    'All interactive elements follow a flat, borderless design language with zero border radius '
    'on buttons and inputs.'
)

add_heading('4.2 Hardware Interfaces', 2)
add_bullet('Camera: Required for product image capture (seller add product flow), accessed via the image_picker package.')
add_bullet('GPS/Location Sensor: Required for rider map tracking, accessed via the geolocator package with user permission.')
add_bullet('Network Adapter: Required for all Supabase communication (Wi-Fi or mobile data).')
add_bullet('Storage: Device internal storage is used by SharedPreferences for session and address data persistence.')
doc.add_paragraph()

add_heading('4.3 Software Interfaces', 2)
add_table(
    ['External System', 'Interface Type', 'Purpose'],
    [
        ('Supabase Auth', 'REST API / JWT', 'User registration, login, OTP verification, password reset, session management.'),
        ('Supabase Database (PostgreSQL)', 'REST API (PostgREST)', 'All data operations: products, orders, users, sellers, riders, notifications.'),
        ('Supabase Storage', 'REST API', 'Upload and retrieve product images from the "products" storage bucket.'),
        ('Google Fonts CDN', 'HTTP', 'Download Inter, Commissioner, and Playfair Display typefaces at runtime.'),
        ('OpenStreetMap', 'HTTP tile server', 'Render map tiles on the Rider Map Screen via flutter_map.'),
        ('Geolocator (OS)', 'Platform plugin', 'Access device GPS for rider real-time location tracking.'),
        ('SharedPreferences (OS)', 'Platform plugin', 'Persist session tokens, delivery addresses, and cart state locally.'),
    ],
    [1.8, 1.5, 3.2]
)

add_heading('4.4 Communication Interfaces', 2)
add_bullet('HTTPS: All communication between the Flutter client and Supabase services uses HTTPS (TLS 1.2 or higher).')
add_bullet('WebSocket (Supabase Realtime): Optional real-time data sync for order status updates (Supabase Realtime channel subscriptions).')
add_bullet('SMTP (Supabase Email): Supabase handles outbound email delivery for OTP codes and account notifications via its configured SMTP provider.')
doc.add_paragraph()

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — NON-FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading('5. Non-Functional Requirements', 1)

add_heading('5.1 Performance Requirements', 2)
add_table(
    ['ID', 'Requirement', 'Target'],
    [
        ('NFR-101', 'Product list loading time on the Landing Screen (first fetch)', 'Under 5 seconds on a standard 4G connection'),
        ('NFR-102', 'Product image loading time per image', 'Under 3 seconds on a standard 4G connection'),
        ('NFR-103', 'Authentication response time (login/signup)', 'Under 3 seconds'),
        ('NFR-104', 'Seller dashboard statistics loading time', 'Under 5 seconds'),
        ('NFR-105', 'Rider map GPS location update frequency', 'At least once every 5 seconds during active navigation'),
        ('NFR-106', 'Product image upload time (single image, < 5MB)', 'Under 10 seconds on a standard 4G connection'),
    ],
    [0.8, 4.0, 1.8]
)

add_heading('5.2 Security Requirements', 2)
add_table(
    ['ID', 'Requirement', 'Details'],
    [
        ('NFR-201', 'Authentication', 'All authentication is handled by Supabase Auth using industry-standard JWT tokens. Passwords are never stored in plain text.'),
        ('NFR-202', 'Row-Level Security (RLS)', 'All Supabase database tables are protected by RLS policies ensuring users can only access records they are authorized to read or modify.'),
        ('NFR-203', 'API Key Protection', 'The Supabase anon key and URL are stored in a .env file excluded from version control.'),
        ('NFR-204', 'Role Enforcement', 'Role-based access control is enforced on both the client (route guards) and server (RLS policies).'),
        ('NFR-205', 'Admin Mobile Block', 'Accounts with role = "admin" accessing the System from a mobile device are redirected to the Unauthorized Screen.'),
        ('NFR-206', 'OTP Expiry', 'OTP codes for email verification and password reset expire after a limited time as enforced by Supabase Auth configuration.'),
    ],
    [0.8, 1.8, 4.0]
)

add_heading('5.3 Usability Requirements', 2)
add_bullet('The application shall support both Android and iOS platforms with a single Flutter codebase.')
add_bullet('The UI shall follow a minimalist design language ensuring intuitive navigation for users with basic smartphone literacy.')
add_bullet('All error messages shall be human-readable and descriptive (e.g., "Incorrect email or password" rather than raw exception codes).')
add_bullet('The application shall display loading indicators (progress spinners) during all asynchronous data operations.')
add_bullet('The application shall display empty-state messages when no data is available (e.g., "No products found", "No orders yet").')
add_bullet('The size guide feature shall be accessible from both the product card and the product detail screen for applicable clothing categories.')
doc.add_paragraph()

add_heading('5.4 Reliability and Availability', 2)
add_bullet('The System shall rely on Supabase\'s managed infrastructure which provides 99.9% uptime SLA for database and authentication services.')
add_bullet('The System shall gracefully handle network failures by displaying retry options rather than crashing.')
add_bullet('The shopping cart shall persist in memory for the duration of the user session and shall survive navigation between screens.')
add_bullet('Authentication sessions shall be restored from Supabase\'s secure session storage on application restart, eliminating the need to log in repeatedly.')
doc.add_paragraph()

add_heading('5.5 Maintainability and Scalability', 2)
add_bullet('The codebase shall follow a clear separation of concerns: screens in /screens, business logic in /services, data models in /models, and reusable widgets in /widgets.')
add_bullet('All Supabase queries shall be encapsulated within service classes to allow future backend changes without modifying screen code.')
add_bullet('The System is designed to support multiple sellers (multi-vendor architecture), with each seller\'s data isolated by seller_id foreign key references.')
add_bullet('The category system is database-driven, allowing new product categories to be added by the Admin without code changes.')
add_bullet('The Supabase PostgreSQL database supports horizontal scaling through connection pooling (PgBouncer) available on Supabase Pro plans.')
doc.add_paragraph()

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — OTHER REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading('6. Other Requirements', 1)

add_heading('6.1 Database Requirements', 2)
add_paragraph(
    'The System uses a PostgreSQL database hosted on Supabase with the following primary tables:'
)
add_table(
    ['Table Name', 'Primary Key', 'Key Fields', 'Purpose'],
    [
        ('users', 'id (BIGINT)', 'email, auth_user_id (UUID), first_name, last_name, phone, role, account_status', 'Stores user profiles for all roles.'),
        ('sellers', 'id (INT)', 'user_id (FK→users), auth_user_id (UUID), store_name, store_slug, address, contact_email, contact_phone, status, commission_rate', 'Stores approved seller profiles and store information.'),
        ('riders', 'id (INT)', 'user_id (FK→users), auth_user_id (UUID), license_number, vehicle_type, address, status', 'Stores approved rider profiles.'),
        ('products', 'id (INT)', 'seller_id (FK→sellers), name, description, price, stock, category_id (FK), approval_status, archive_status, delivery_options, condition', 'Stores all product listings.'),
        ('product_images', 'id (INT)', 'product_id (FK→products), image_url', 'Stores image URLs for product listings.'),
        ('product_variants', 'id (INT)', 'product_id (FK→products), color, size, stock', 'Stores size and color variants with individual stock levels.'),
        ('categories', 'id (INT)', 'name (UNIQUE)', 'Lookup table for product categories.'),
        ('orders', 'id (INT/BIGINT)', 'seller_id, buyer_id, product_id, total_amount, status, created_at', 'Stores order records linking buyers, sellers, and products.'),
        ('notifications', 'id (INT)', 'seller_id (FK→sellers), type, title, message, product_id, is_read, created_at', 'Stores seller notifications for product approval events.'),
        ('seller_daily_stats', 'id', 'seller_id, date, sales_amount, orders_count', 'Stores daily revenue data for the seller dashboard chart.'),
        ('inventory_transactions', 'id', 'product_id, seller_id, quantity_change, reason, created_at', 'Audit log for stock quantity changes.'),
    ],
    [1.2, 1.3, 2.8, 1.7]
)
add_paragraph(
    'All tables are protected by Supabase Row-Level Security (RLS) policies. Public read '
    'access is granted on the sellers table to allow unauthenticated buyers to view store '
    'information. All write operations require authentication and are scoped to the '
    'authenticated user\'s own records.'
)

add_heading('6.2 Legal and Compliance Requirements', 2)
add_bullet('The System shall collect only the minimum personal data necessary for operation (email, name, phone, address).')
add_bullet('User passwords shall never be stored in plain text; authentication is delegated entirely to Supabase Auth which uses bcrypt hashing.')
add_bullet('Seller and rider applicants implicitly consent to their personal information being reviewed by the platform administrator as part of the approval process.')
add_bullet('Product images uploaded by sellers are stored in a publicly accessible Supabase Storage bucket; sellers are responsible for ensuring uploaded images do not violate copyright or intellectual property rights.')
add_bullet('The System does not currently process real financial transactions. Future payment integration must comply with applicable financial regulations including the Bangko Sentral ng Pilipinas (BSP) guidelines for electronic payment systems in the Philippines.')
add_bullet('The application complies with Google Play Store and Apple App Store data privacy guidelines regarding location access, which is requested only when the Rider Map feature is actively used.')
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

output_path = r'c:\Users\Admin\OneDrive\Documents\GitHub\E-Commerce_Mobile\Varon_SRS.docx'
doc.save(output_path)
print(f'SRS document saved to: {output_path}')
